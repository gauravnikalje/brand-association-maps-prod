import pandas as pd
from docx import Document
import os

BASE = r"C:\Users\Administrator\Desktop\BAM\Brand Association Maps"
OUT  = r"C:\Users\Administrator\Desktop\BAM\BAM_context.md"

lines = []

def h(title):
    lines.append(f"\n\n---\n## {title}\n")

def p(text):
    lines.append(str(text))

# ──────────────────────────────────────────────
# 1. BAM Execution Process Steps.docx
# ──────────────────────────────────────────────
h("BAM Execution Process Steps (docx)")
try:
    doc = Document(os.path.join(BASE, "BAM Execution Process Steps.docx"))
    for para in doc.paragraphs:
        if para.text.strip():
            p(para.text)
    # Also grab tables
    for table in doc.tables:
        for row in table.rows:
            p(" | ".join(cell.text.strip() for cell in row.cells))
    p("\n[docx READ OK]")
except Exception as e:
    p(f"[ERROR reading docx: {e}]")

# ──────────────────────────────────────────────
# 2. Input_Data_1.xlsx
# ──────────────────────────────────────────────
h("Input_Data_1.xlsx")
try:
    xl = pd.ExcelFile(os.path.join(BASE, "Input_Data_1.xlsx"))
    p(f"Sheets: {xl.sheet_names}")
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        p(f"\n### Sheet: {sheet}")
        p(f"Shape: {df.shape}")
        p(f"Columns: {list(df.columns)}")
        p(f"Dtypes:\n{df.dtypes.to_string()}")
        p(f"\nFirst 3 rows:\n{df.head(3).to_string()}")
        p(f"\nNull counts:\n{df.isnull().sum().to_string()}")
        # If there's a text/message column, show unique value counts
        for col in df.columns:
            if df[col].dtype == object:
                p(f"\nUnique values in '{col}' (top 10): {df[col].value_counts().head(10).to_dict()}")
    p("[Input_Data_1 READ OK]")
except Exception as e:
    p(f"[ERROR: {e}]")

# ──────────────────────────────────────────────
# 3. Input_Data_2.xlsx
# ──────────────────────────────────────────────
h("Input_Data_2.xlsx")
try:
    xl = pd.ExcelFile(os.path.join(BASE, "Input_Data_2.xlsx"))
    p(f"Sheets: {xl.sheet_names}")
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        p(f"\n### Sheet: {sheet}")
        p(f"Shape: {df.shape}")
        p(f"Columns: {list(df.columns)}")
        p(f"\nFirst 5 rows:\n{df.head(5).to_string()}")
        for col in df.columns:
            if df[col].dtype == object:
                p(f"\nUnique values in '{col}' (top 10): {df[col].value_counts().head(10).to_dict()}")
    p("[Input_Data_2 READ OK]")
except Exception as e:
    p(f"[ERROR: {e}]")

# ──────────────────────────────────────────────
# 4. Bigram tagging_taxonomy.xlsx
# ──────────────────────────────────────────────
h("Bigram tagging_taxonomy.xlsx")
try:
    xl = pd.ExcelFile(os.path.join(BASE, "Bigram tagging_taxonomy.xlsx"))
    p(f"Sheets: {xl.sheet_names}")
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        p(f"\n### Sheet: {sheet}")
        p(f"Shape: {df.shape}")
        p(f"Columns: {list(df.columns)}")
        p(f"\nFirst 10 rows:\n{df.head(10).to_string()}")
        # Show unique T1/T2/T3 values
        for col in df.columns:
            if 'attribute' in str(col).lower() or 'T1' in str(col) or 'T2' in str(col) or 'T3' in str(col):
                p(f"\nUnique '{col}': {sorted(df[col].dropna().unique().tolist())}")
    p("[Bigram taxonomy READ OK]")
except Exception as e:
    p(f"[ERROR: {e}]")

# ──────────────────────────────────────────────
# 5. monogram tagging_taxonomy.xlsx
# ──────────────────────────────────────────────
h("monogram tagging_taxonomy.xlsx")
try:
    xl = pd.ExcelFile(os.path.join(BASE, "monogram tagging_taxonomy.xlsx"))
    p(f"Sheets: {xl.sheet_names}")
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        p(f"\n### Sheet: {sheet}")
        p(f"Shape: {df.shape}")
        p(f"Columns: {list(df.columns)}")
        p(f"\nFirst 10 rows:\n{df.head(10).to_string()}")
        for col in df.columns:
            if 'attribute' in str(col).lower():
                p(f"\nUnique '{col}': {sorted(df[col].dropna().unique().tolist())}")
    p("[Monogram taxonomy READ OK]")
except Exception as e:
    p(f"[ERROR: {e}]")

# ──────────────────────────────────────────────
# 6. Rawdata_ref.xlsx  (large — only headers + 5 rows + shape)
# ──────────────────────────────────────────────
h("Rawdata_ref.xlsx (large file — headers + sample only)")
try:
    xl = pd.ExcelFile(os.path.join(BASE, "Rawdata_ref.xlsx"))
    p(f"Sheets: {xl.sheet_names}")
    for sheet in xl.sheet_names:
        df = xl.parse(sheet, nrows=5)
        p(f"\n### Sheet: {sheet}")
        p(f"Columns: {list(df.columns)}")
        p(f"\nFirst 5 rows:\n{df.head(5).to_string()}")
        # get full row count without loading all data
        df_full = xl.parse(sheet, usecols=[0])
        p(f"Approx total rows: {len(df_full)}")
        for col in df.columns:
            if df[col].dtype == object:
                p(f"\nSample unique values in '{col}': {df[col].dropna().unique()[:5].tolist()}")
    p("[Rawdata_ref READ OK]")
except Exception as e:
    p(f"[ERROR: {e}]")

# ──────────────────────────────────────────────
# Write out
# ──────────────────────────────────────────────
with open(OUT, "w", encoding="utf-8") as f:
    f.write("# BAM — Brand Association Maps: Full Context\n")
    f.write("_Auto-generated by read_context.py — do not edit manually_\n")
    f.write("\n".join(lines))

print(f"\n✅ Context saved to: {OUT}")
